"""
word_generator.py
灵活可调用的Word文档生成器，支持完全通过接口/JSON配置控制样式
核心支持：一/二/三/四级标题、段落、表格样式的全自定义
修改后特性：支持直接接收str格式的JSON样式配置字符串
"""
from io import BytesIO
import json
import os
import requests
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn as oxml_qn
import re

# 配置常量
# 获取脚本自身的目录（不是运行时的工作目录）
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_CONFIG_PATH = os.path.join(SCRIPT_DIR, "default_word_styles.json")
TIMEOUT_SECONDS = 15  # 接口超时时间
# 配置文件编码
CONFIG_ENCODING = "utf-8"

class CustomWordGenerator:
    """可完全通过JSON配置控制样式的Word文档生成器（支持字符串JSON配置）"""
    
    def __init__(self, 
                 config_json_str: Optional[str] = None,  # 新增：JSON字符串配置
                 config_api_url: Optional[str] = None,
                 api_headers: Optional[Dict] = None,
                 api_params: Optional[Dict] = None):
        """
        初始化Word生成器
        
        Args:
            config_json_str: JSON格式的样式配置字符串（优先级最高）
            config_api_url: 样式配置接口地址
            api_headers: 接口请求头（如token、Content-Type）
            api_params: 接口请求参数（GET/POST）
        """
        self.doc = None
        self.config_json_str = config_json_str  # 存储JSON字符串
        self.config_api_url = config_api_url
        self.api_headers = api_headers or {"Content-Type": "application/json"}
        self.api_params = api_params or {}
        
        # 加载配置（优先级：JSON字符串 > 接口 > 内置默认）
        self.styles_config = self._load_styles_config()
        
    def _get_default_config(self) -> Dict:
        """从JSON文件加载默认样式配置"""
        default_config = {}
         # 1. 检查文件是否存在
        if not os.path.exists(DEFAULT_CONFIG_PATH):
            raise FileNotFoundError(
                f"配置文件不存在！当前路径：{os.getcwd()}\n"
                f"请确认 {DEFAULT_CONFIG_PATH} 在该目录下"
            )
        try:
            # 读取默认配置文件
            with open(DEFAULT_CONFIG_PATH, "r", encoding=CONFIG_ENCODING) as f:
                default_config = json.load(f)
            print(f"✅ 成功加载默认样式配置文件：{DEFAULT_CONFIG_PATH}")
        except FileNotFoundError:
            print(f"⚠️ 默认配置文件 {DEFAULT_CONFIG_PATH} 不存在，使用内置极简配置兜底")
            # 极简兜底配置（防止完全加载失败）
            default_config = {
                "document": {"page_size": "A4", "orientation": "portrait"},
                "styles": {
                    "normal": {"font": {"name": "微软雅黑", "size": 12}, "paragraph": {"alignment": "LEFT"}},
                    "heading1": {"font": {"name": "微软雅黑", "size": 24, "bold": True}}
                }
            }
        except json.JSONDecodeError as e:
            print(f"⚠️ 默认配置文件解析失败：{str(e)}，使用内置极简配置兜底")
            default_config = {
                "document": {"page_size": "A4", "orientation": "portrait"},
                "styles": {
                    "normal": {"font": {"name": "微软雅黑", "size": 12}, "paragraph": {"alignment": "LEFT"}},
                    "heading1": {"font": {"name": "微软雅黑", "size": 24, "bold": True}}
                }
            }
        return default_config
    
    def _load_from_json_str(self) -> Optional[Dict]:
        """从JSON字符串加载样式配置"""
        if not self.config_json_str:
            return None
        
        try:
            json_config = json.loads(self.config_json_str)
            print(f"✅ 成功从JSON字符串加载样式配置")
            return json_config
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON字符串解析失败：{str(e)}，降级到接口/默认配置")
            return None
    
    def _load_from_api(self) -> Optional[Dict]:
        """从接口加载样式配置（支持GET/POST）"""
        if not self.config_api_url:
            return None
        
        try:
            # 自动判断请求方式（根据Content-Type）
            if self.api_headers.get("Content-Type") == "application/json":
                response = requests.post(
                    url=self.config_api_url,
                    headers=self.api_headers,
                    json=self.api_params,
                    timeout=TIMEOUT_SECONDS
                )
            else:
                response = requests.get(
                    url=self.config_api_url,
                    headers=self.api_headers,
                    params=self.api_params,
                    timeout=TIMEOUT_SECONDS
                )
            
            response.raise_for_status()
            api_config = response.json()
            print(f"✅ 成功从接口加载样式配置：{self.config_api_url}")
            return api_config
        
        except requests.exceptions.RequestException as e:
            print(f"⚠️ 接口加载配置失败：{str(e)}，降级到默认配置")
            return None
    
    def _convert_units(self, config: Dict) -> Dict:
        """将JSON中的数值转换为docx单位对象"""
        # 页面边距（Cm）
        if "document" in config:
            doc_config = config["document"]  # 定义变量
            if "margins" in doc_config:  # 先判空 margins 是否存在
                margins = doc_config["margins"]
                for k, v in margins.items():
                    margins[k] = Cm(v)
        
        # 字体大小/段落间距（Pt）
        def _convert_font_para(style_dict):
            if "font" in style_dict and "size" in style_dict["font"]:
                style_dict["font"]["size"] = Pt(style_dict["font"]["size"])
            if "paragraph" in style_dict:
                para = style_dict["paragraph"]
                for k in ["space_before", "space_after", "left_indent", "right_indent", 
                          "first_line_indent", "hanging_indent"]:
                    if k in para:
                        para[k] = Pt(para[k])
        
        # 处理所有样式的单位
        if "styles" in config:
            for style in config["styles"].values():
                _convert_font_para(style)
        
        # 表格样式单位
        if "table" in config:
            table_config = config["table"]
            # 表头字体
            if "header" in config["table"] and "font" in config["table"]["header"]and "size" in table_config["header"]["font"]:
                config["table"]["header"]["font"]["size"] = Pt(config["table"]["header"]["font"]["size"])
                if "header" in table_config:
                    # 表头段落间距
                    for k in ["space_before", "space_after"]:
                        if k in config["table"]["header"]:
                            config["table"]["header"][k] = Pt(config["table"]["header"][k])
            # 行字体/行高
            if "row" in config["table"]:
                row_config = table_config["row"]
                if "font" in row_config and "size" in row_config["font"]:  # size判空
                    row_config["font"]["size"] = Pt(row_config["font"]["size"])
                if "height" in config["table"]["row"]:
                    config["table"]["row"]["height"] = Pt(config["table"]["row"]["height"])
                # 行高规则转换
                if "height_rule" in config["table"]["row"]:
                    rule_map = {"AUTO": WD_ROW_HEIGHT_RULE.AUTO, "FIXED": WD_ROW_HEIGHT_RULE.EXACTLY}
                    config["table"]["row"]["height_rule"] = rule_map.get(
                        config["table"]["row"]["height_rule"].upper(), WD_ROW_HEIGHT_RULE.AUTO
                    )
            # 列宽转换
            if "column_width" in config["table"]:
                for k, v in config["table"]["column_width"].items():
                    config["table"]["column_width"][k] = Cm(v)
        
        # 页眉页脚字体大小
        if "header_footer" in config:
            hf_config = config["header_footer"]  # 变量简化读取
            for part in ["header", "footer"]:
                if part in hf_config and "font_size" in hf_config[part]:
                    hf_config[part]["font_size"] = Pt(hf_config[part]["font_size"])
        return config
    
    def _merge_configs(self, default: Dict, user_config: Dict) -> Dict:
        """递归合并配置（用户配置覆盖默认）"""
        for key, value in user_config.items():
            if key in default and isinstance(default[key], dict) and isinstance(value, dict):
                default[key] = self._merge_configs(default[key], value)
            else:
                default[key] = value
        return default
    
    def _load_styles_config(self) -> Dict:
        """加载最终样式配置"""
        # 1. JSON字符串配置（优先级最高）
        json_str_config = self._load_from_json_str()
        if json_str_config:
            merged = self._merge_configs(self._get_default_config(), json_str_config)
            return self._convert_units(merged)
        
        # 2. 接口配置
        api_config = self._load_from_api()
        if api_config:
            merged = self._merge_configs(self._get_default_config(), api_config)
            return self._convert_units(merged)
        
        # 3. 默认配置
        return self._convert_units(self._get_default_config())
    
    def create_document(self) -> Document:
        """创建新文档并应用基础配置"""
        self.doc = Document()
        self._setup_page_settings()
        self._create_custom_styles()
        return self.doc
    
    def _setup_page_settings(self):
        """应用页面配置"""
        if not self.doc:
            raise ValueError("文档未初始化，请先调用create_document()")
        
        section = self.doc.sections[0]
        doc_config = self.styles_config["document"]
        
        # 页面边距
        margins = doc_config["margins"]
        section.top_margin = margins["top"]
        section.bottom_margin = margins["bottom"]
        section.left_margin = margins["left"]
        section.right_margin = margins["right"]
        section.header_distance = margins["header"]
        section.footer_distance = margins["footer"]
        
        # 页面方向
        if doc_config["orientation"].upper() == "LANDSCAPE":
            section.orientation = WD_SECTION.LANDSCAPE
    
    def _create_custom_styles(self):
        """根据配置创建自定义样式（核心：标题/段落样式全自定义）"""
        styles_config = self.styles_config["styles"]
        
        for style_key, style_config in styles_config.items():
            # 标题样式（使用内置Heading样式）
            if style_key.startswith("heading"):
                level = int(style_key.replace("heading", ""))
                if 1 <= level <= 4:
                    style_name = f"Heading {level}"
                    style = self.doc.styles[style_name]
                else:
                    continue
            # 文档标题
            elif style_key == "title":
                style = self.doc.styles["Title"]
            # 自定义段落样式
            else:
                style_id = f"Custom{style_key.capitalize()}"
                # 如果样式已存在，先删除
                if style_id in [s.name for s in self.doc.styles]:
                    self.doc.styles[self.doc.styles.index(style_id)].delete()
                style = self.doc.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = self.doc.styles["Normal"]
            
            # 应用样式配置
            self._apply_style_config(style, style_config)
    
    def _apply_style_config(self, style, config: Dict):
        """应用样式配置到docx样式对象"""
        # 字体配置
        font_config = config.get("font", {})
        if "name" in font_config:
            style.font.name = font_config["name"]
            # 设置中文字体
            try:
                r_fonts = style._element.rPr.rFonts
                r_fonts.set(qn("w:eastAsia"), font_config["name"])
                r_fonts.set(qn("w:ascii"), font_config["name"])
            except:
                pass
        
        for attr in ["size", "bold", "italic", "underline"]:
            if attr in font_config:
                setattr(style.font, attr, font_config[attr])
        
        if "color" in font_config:
            style.font.color.rgb = RGBColor.from_string(font_config["color"])
        
        # 段落配置
        para_config = config.get("paragraph", {})
        if "alignment" in para_config:
            align_map = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            style.paragraph_format.alignment = align_map.get(
                para_config["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.LEFT
            )
        
        for attr in ["space_before", "space_after", "left_indent", "right_indent",
                     "first_line_indent", "hanging_indent", "line_spacing",
                     "keep_with_next", "page_break_before"]:
            if attr in para_config:
                setattr(style.paragraph_format, attr, para_config[attr])
        
        # 大纲级别
        if "outline_level" in config:
            ppr = style._element.get_or_add_pPr()
            outline_lvl = OxmlElement('w:outlineLvl')
            outline_lvl.set(qn('w:val'), str(config["outline_level"] - 1))
            ppr.append(outline_lvl)
    
    # ===================== 核心功能：标题/段落/表格添加 =====================
    def add_heading(self, text: str, level: int = 1, style_override: Optional[Dict] = None):
        """
        添加标题（完全受接口JSON控制）
        Args:
            text: 标题文本
            level: 1-4级标题
            style_override: 临时样式覆盖（优先级高于配置）
        """
        if not self.doc:
            raise ValueError("文档未初始化")
        
        # 校验级别
        level = max(1, min(4, level))
        style_key = f"heading{level}"
        
        # 获取基础配置
        base_config = self.styles_config["styles"].get(style_key, {})
        if style_override:
            base_config = self._merge_configs(base_config, style_override)
        
        # 创建标题段落
        para = self.doc.add_paragraph(text, style=f"Heading {level}")
        para.clear()
        run = para.add_run(text)
        
        # 应用字体样式
        self._apply_run_style(run, base_config.get("font", {}))
        
        # 应用段落样式
        para_config = base_config.get("paragraph", {})
        if "alignment" in para_config:
            align_map = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT
            }
            para.alignment = align_map.get(para_config["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.LEFT)
        
        for attr in ["space_before", "space_after", "keep_with_next", "page_break_before"]:
            if attr in para_config:
                setattr(para.paragraph_format, attr, para_config[attr])
    
    def add_paragraph(self, text: str, style: str = "normal", style_override: Optional[Dict] = None):
        """
        添加段落（完全受接口JSON控制）
        Args:
            text: 段落文本
            style: 样式名（normal/quote/list_item等）
            style_override: 临时样式覆盖
        """
        if not self.doc:
            raise ValueError("文档未初始化")
        
        # 获取基础样式配置
        base_config = self.styles_config["styles"].get(style, self.styles_config["styles"]["normal"])
        if style_override:
            base_config = self._merge_configs(base_config, style_override)
        
        # 创建段落
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        # 应用字体样式
        self._apply_run_style(run, base_config.get("font", {}))
        
        # 应用段落样式
        para_config = base_config.get("paragraph", {})
        if "alignment" in para_config:
            align_map = {
                "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            }
            para.alignment = align_map.get(para_config["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
        
        for attr in ["space_before", "space_after", "left_indent", "right_indent",
                     "first_line_indent", "hanging_indent", "line_spacing"]:
            if attr in para_config:
                setattr(para.paragraph_format, attr, para_config[attr])
    
    def add_table(self, data: List[List[Any]], headers: Optional[List[str]] = None, title: Optional[str] = None):
        """
        添加表格（完全受接口JSON控制）
        Args:
            data: 表格数据
            headers: 表头
            title: 表格标题
        """
        if not self.doc or not data:
            return
        
        table_config = self.styles_config["table"]
        col_width = table_config["column_width"].get("default", Cm(6))
        
        # 添加表格标题
        if title:
            self.add_paragraph(title, style="normal", style_override={
                "font": {"bold": True, "size": 16},
                "paragraph": {"space_before": 12, "space_after": 6}
            })
        
        # 计算行列数
        num_rows = len(data) + (1 if headers else 0)
        num_cols = max(len(row) for row in data) if data else 0
        if headers:
            num_cols = max(num_cols, len(headers))
        
        # 创建表格
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = table_config["default_style"]
        table.autofit = False
        
        # 设置列宽
        for col in table.columns:
            col.width = col_width
        
        # 填充表头
        if headers:
            header_config = table_config["header"]
            header_row = table.rows[0]
            
            # 表头行高
            header_row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
            header_row.height = Pt(30)
            
            for i, header_text in enumerate(headers):
                if i >= num_cols:
                    break
                cell = header_row.cells[i]
                cell.text = str(header_text)
                
                # 表头样式
                for para in cell.paragraphs:
                    # 清空原有内容
                    para.clear()
                    run = para.add_run(str(header_text))
                    
                    # 应用表头字体样式
                    self._apply_run_style(run, header_config["font"])
                    
                    # 表头对齐
                    align_map = {
                        "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                        "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                        "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT
                    }
                    para.alignment = align_map.get(header_config["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.CENTER)
                    
                    # 表头段落间距
                    para.paragraph_format.space_before = header_config.get("space_before", Pt(2))
                    para.paragraph_format.space_after = header_config.get("space_after", Pt(2))
                
                # 表头背景色
                self._set_cell_background(cell, header_config["background_color"])
        
        # 填充数据行
        row_config = table_config["row"]
        start_row = 1 if headers else 0
        
        for row_idx in range(start_row, num_rows):
            data_row = table.rows[row_idx]
            data_idx = row_idx - start_row
            
            # 行高配置
            data_row.height_rule = row_config["height_rule"]
            data_row.height = row_config["height"]
            
            # 行背景色（奇偶行）
            bg_color = row_config["background_even"] if row_idx % 2 == 0 else row_config["background_odd"]
            
            for col_idx in range(num_cols):
                cell = data_row.cells[col_idx]
                cell_text = data[data_idx][col_idx] if data_idx < len(data) and col_idx < len(data[data_idx]) else ""
                cell.text = str(cell_text)
                
                # 行样式
                for para in cell.paragraphs:
                    para.clear()
                    run = para.add_run(str(cell_text))
                    
                    # 应用行字体样式
                    self._apply_run_style(run, row_config["font"])
                    
                    # 行对齐
                    align_map = {
                        "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
                        "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
                        "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT
                    }
                    para.alignment = align_map.get(row_config["alignment"].upper(), WD_PARAGRAPH_ALIGNMENT.CENTER)
                
                # 行背景色
                self._set_cell_background(cell, bg_color)
        
        # 表格居中
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.doc.add_paragraph()
    
    # ===================== 辅助方法 =====================
    def _apply_run_style(self, run, font_config: Dict):
        """应用字符样式"""
        if "name" in font_config:
            run.font.name = font_config["name"]
            try:
                r_fonts = run.element.rPr.rFonts
                r_fonts.set(qn("w:eastAsia"), font_config["name"])
                r_fonts.set(qn("w:ascii"), font_config["name"])
            except:
                pass
        
        for attr in ["size", "bold", "italic", "underline"]:
            if attr in font_config:
                setattr(run.font, attr, font_config[attr])
        
        if "color" in font_config:
            run.font.color.rgb = RGBColor.from_string(font_config["color"])
    
    def _set_cell_background(self, cell, color_hex: str):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _fix_escape_characters(self, text: str) -> str:
        """修复转义字符"""
        # 1. 去掉 ** 和 &nbsp;&nbsp;
        text = text.replace('**', '').replace('&nbsp;&nbsp;', '')  
        # 2. 给 1、2、3、 前面加换行
        text = re.sub(r'(\d+、)', r'\n\1', text)
        # 3. 给 (a)、(b)、(c)、 前面加换行（支持中英文括号） 
        text = re.sub(r'(?<!\n)(（[ａ-ｚA-Z]）、?)', r'\n\n\1', text)
        # 4. 基础换行清理（保留你原来的逻辑）
        text = text.replace('\\n', '\n').replace('  \n', '\n').replace('<br>', '\n')
        # text = re.sub(r'^-\([a-zA-Z0-9]\)\s*', '', text)
        # text = re.sub(r'([（(]\d+[）)])、', r'\n\1、', text)
        # 5. 多余空行清理（最终整洁）
        # text = re.sub(r'\n+', '\n', text).strip()
        return text.strip()
    
    def add_markdown_content(self, markdown_text: str):
        """添加Markdown内容（样式受配置控制）"""
        markdown_text = self._fix_escape_characters(markdown_text)
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            if not line:
                self.doc.add_paragraph()
                i += 1
                continue
            
            # 标题
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                self.add_heading(heading_match.group(2), level=level)
                i += 1
                continue
            
            # 列表
            if re.match(r'^([\-\+]|\(\d+\))\s+.+$', line):
                items = []
                while i < len(lines) and re.match(r'^([\-\+]|\(\d+\))\s+.+$', lines[i]):
                    items.append(re.sub(r'^([\-\+]|\(\d+\))\s+', '', lines[i]))
                    i += 1
                for item in items:
                    self.add_paragraph(f"·{item}", style="list_item")
                continue
            
            # 表格
            if '|' in line and re.match(r'^\|.*\|$', line):
                table_lines = []
                while i < len(lines) and '|' in lines[i] and re.match(r'^\|.*\|$', lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                if len(table_lines) >= 2:
                    headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                    data_start = 2 if '---' in table_lines[1] else 1
                    data = []
                    for line in table_lines[data_start:]:
                        data.append([cell.strip() for cell in line.split('|')[1:-1]])
                    self.add_table(data, headers=headers)
                continue
            
            # 普通段落
            para_text = line
            i += 1
            while i < len(lines) and lines[i].strip() and not re.match(r'^[#\*\-\+|]', lines[i]):
                para_text += ' ' + lines[i].strip()
                i += 1
            self.add_paragraph(para_text)
    
    def save_to_stream(self, stream: BytesIO) -> bool:
        """
        保存文档到字节流（内存）
        
        Args:
            stream: 字节流对象
            
        Returns:
            bool: 是否保存成功
        """
        if not self.doc:
            raise ValueError("文档未初始化，请先调用 create_document()")
        
        try:
            self.doc.save(stream)
            return True
        except Exception as e:
            print(f"保存到流失败: {str(e)}")
            return False
    
    def save_to_bytes(self) -> Optional[bytes]:
        """
        保存文档到字节数据（直接返回 bytes）
        
        Returns:
            bytes: Word文档的字节数据，失败返回None
        """
        stream = BytesIO()
        if self.save_to_stream(stream):
            stream.seek(0)  # 重置指针到开头
            return stream.getvalue()
        return None
    
    def save(self, filepath: str) -> Optional[str]:
        """保存文档"""
        if not self.doc:
            return None
        # 空路径校验 + 兜底
        if not filepath or filepath.strip() == "":
            # 生成临时兜底路径
            timestamp = int(time.time())
            filepath = f"report_fallback_{timestamp}.docx"
            printf("传入空文件路径，自动使用兜底路径")
        try:
            # 确保目录存在（处理绝对/相对路径）
            dir_path = os.path.dirname(filepath)
            if dir_path:  # 仅当目录非空时创建（避免创建空目录）
                os.makedirs(dir_path, exist_ok=True)
            self.doc.save(filepath)
           
            return filepath
        except Exception as e:
            printf(" 保存文档失败")
            return None



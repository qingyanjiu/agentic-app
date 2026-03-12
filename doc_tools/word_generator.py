"""
word_generator.py
灵活可调用的Word文档生成器，支持自定义样式
"""

import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional, Union
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import qn as oxml_qn
import re

class CustomWordGenerator:
    """可自定义样式的Word文档生成器"""
    
    def __init__(self):
        """
        初始化Word生成器
        
        Args:
            config_path: 样式配置文件路径，如果为None则使用默认配置
        """
        self.doc = None
        self.styles_config = self._get_default_config()
        self.current_heading_level = 0
        self.table_data = []
        self.in_table = False
        
    def _get_default_config(self) -> Dict:
        """获取默认配置"""
        return {
            "document": {
                "page_size": "A4",
                "orientation": "portrait",
                "margins": {
                    "top": Cm(2.54),
                    "bottom": Cm(2.54),
                    "left": Cm(3.17),
                    "right": Cm(3.17),
                    "header": Cm(1.27),
                    "footer": Cm(1.27)
                }
            },
            "styles": {
                "title": {
                    "name": "文档标题",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(22),
                        "bold": True,
                        "color": "1F4E78",
                        "underline": False
                    },
                    "paragraph": {
                        "alignment": "CENTER",
                        "space_before": Pt(0),
                        "space_after": Pt(24),
                        "line_spacing": 1.5
                    }
                },
                "heading1": {
                    "name": "一级标题",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(26),
                        "bold": True,
                        "italic": False,
                        "color": "000000"
                    },
                    "paragraph": {
                        "alignment": "LEFT",
                        #标题段落上方的间距
                        "space_before": Pt(12),
                        #标题段落下方的间距
                        "space_after": Pt(6),
                        "keep_with_next": False
                    },
                    "outline_level": 1
                },
                "heading2": {
                    "name": "二级标题",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(18),
                        "bold": True,
                        "italic": False,
                        "color": "2E74B5"
                    },
                    "paragraph": {
                        "alignment": "LEFT",
                        "space_before": Pt(19),
                        "space_after": Pt(7),
                        "keep_with_next": True,
                         "page_break_before": True  # 关键：段前分页，与上一段不同页
                    },
                    "outline_level": 2
                },
                "heading3": {
                    "name": "三级标题",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(14),
                        "bold": True,
                        "italic": False,
                        "color": "000000"
                    },
                    "paragraph": {
                        "alignment": "LEFT",
                        "space_before": Pt(16),
                        "space_after": Pt(6),
                        "keep_with_next": True
                    },
                    "outline_level": 3
                },
                "heading4": {
                    "name": "四级标题",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(11),
                        "bold": True,
                        "italic": False,
                        "color": "000000"
                    },
                    "paragraph": {
                        "alignment": "LEFT",
                        "space_before": Pt(16),
                        "space_after": Pt(6),
                        "keep_with_next": True
                    },
                    "outline_level": 4
                },
                "normal": {
                    "name": "正文",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(12),
                        "bold": False,
                        "color": "000000"
                    },
                    "paragraph": {
                        "alignment": "JUSTIFY",
                        "space_before": Pt(6),
                        "space_after": Pt(6),
                        "first_line_indent": Pt(21),
                        "line_spacing": 1.5
                    }
                },
                "quote": {
                    "name": "引用",
                    "font": {
                        "name": "楷体",
                        "size": Pt(11),
                        "italic": True,
                        "color": "666666"
                    },
                    "paragraph": {
                        "alignment": "JUSTIFY",
                        "left_indent": Pt(21),
                        "right_indent": Pt(21),
                        "space_before": Pt(6),
                        "space_after": Pt(6)
                    }
                },
                "list_item": {
                    "name": "列表项",
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(12)
                    },
                    "paragraph": {
                        "alignment": "LEFT",
                        "left_indent": Pt(21),
                        "hanging_indent": Pt(21),
                        "space_before": Pt(0),
                        "space_after": Pt(0)
                    }
                }
            },
            "table": {
                "default_style": "Light Grid Accent 1",
                "header": {
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(14),
                        "bold": True,
                        "color": "FFFFFF"
                    },
                    "background_color": "2E74B5",
                    "alignment": "LEFT"
                },
                "row": {
                    "font": {
                        "name": "微软雅黑",
                        "size": Pt(10)
                    },
                    "background_even": "FFFFFF",
                    "background_odd": "F2F2F2",
                    "alignment": "LEFT",
                    "height": Pt(30)
                }
            },
            "toc": {
                "enabled": True,
                "title": "目录",
                "title_style": "heading1",
                "levels": [1, 2, 3],
                "leader": "dot",
                "tab_stops": [
                    {"position": Cm(14), "alignment": "RIGHT"}
                ]
            },
            "header_footer": {
                "enabled": True,
                "header": {
                    "text": "{report_title}",
                    "font_size": Pt(9),
                    "alignment": "CENTER"
                },
                "footer": {
                    "left": "生成时间: {timestamp}",
                    "center": "第 {page_number} 页，共 {total_pages} 页",
                    "right": "园区数字化运营中心",
                    "font_size": Pt(9)
                }
            }
        }
    
    def _load_config(self) -> Dict:
        """加载配置文件"""
        return self._get_default_config()
    
    def _merge_configs(self, default: Dict, user_config: Dict) -> Dict:
        """递归合并配置"""
        for key, value in user_config.items():
            if key in default and isinstance(default[key], dict) and isinstance(value, dict):
                default[key] = self._merge_configs(default[key], value)
            else:
                default[key] = value
        return default
    
    def create_document(self) -> Document:
        """创建新文档"""
        self.doc = Document()
        self._setup_page_settings()
        self._create_custom_styles()
        return self.doc
    
    def _setup_page_settings(self):
        """设置页面设置"""
        sections = self.doc.sections
        for section in sections:
            # 页面边距
            config_margins = self.styles_config["document"]["margins"]
            section.top_margin = config_margins["top"]
            section.bottom_margin = config_margins["bottom"]
            section.left_margin = config_margins["left"]
            section.right_margin = config_margins["right"]
            section.header_distance = config_margins["header"]
            section.footer_distance = config_margins["footer"]
            
            # 页面方向
            if self.styles_config["document"]["orientation"] == "landscape":
                section.orientation = WD_SECTION.LANDSCAPE
    
    def _create_custom_styles(self):
        """创建自定义样式"""
        styles_config = self.styles_config["styles"]
        
        # 为每种样式创建自定义样式
        for style_name, style_config in styles_config.items():
            if style_name == "title":
                # 文档标题样式
                style = self.doc.styles["Title"]
            elif style_name.startswith("heading"):
                # 标题样式（使用内置样式）
                level = int(style_name.replace("heading", ""))
                style_name_internal = f"Heading {level}"
                style = self.doc.styles[style_name_internal]
            else:
                # 创建自定义样式
                style_id = f"Custom{style_name.capitalize()}"
                style = self.doc.styles.add_style(style_id, WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = self.doc.styles["Normal"]
            
            # 应用样式配置
            self._apply_style_config(style, style_config)
    
    def _apply_style_config(self, style, config: Dict):
        """应用样式配置到样式对象"""
        # 字体设置
        font_config = config.get("font", {})
        if "name" in font_config:
            style.font.name = font_config["name"]
            # 设置中文字体
            try:
                r_fonts = style._element.rPr.rFonts
                r_fonts.set(qn("w:eastAsia"), font_config["name"])
                r_fonts.set(qn("w:ascii"), font_config["name"])
            except:
                pass
        
        if "size" in font_config:
            style.font.size = font_config["size"]
        if "bold" in font_config:
            style.font.bold = font_config["bold"]
        if "italic" in font_config:
            style.font.italic = font_config["italic"]
        if "underline" in font_config:
            style.font.underline = font_config["underline"]
        if "color" in font_config:
            style.font.color.rgb = RGBColor.from_string(font_config["color"])
        
        # 段落设置
        para_config = config.get("paragraph", {})
        if "alignment" in para_config:
            style.paragraph_format.alignment = self._get_alignment(para_config["alignment"])
        
        for key in ["space_before", "space_after", "left_indent", "right_indent", 
                   "first_line_indent", "hanging_indent"]:
            if key in para_config:
                setattr(style.paragraph_format, key, para_config[key])
        
        if "line_spacing" in para_config:
            style.paragraph_format.line_spacing = para_config["line_spacing"]
        
        if "keep_with_next" in para_config:
            style.paragraph_format.keep_with_next = para_config["keep_with_next"]
        
        # 大纲级别
        if "outline_level" in config:
            ppr = style._element.get_or_add_pPr()
            outline_lvl = OxmlElement('w:outlineLvl')
            outline_lvl.set(qn('w:val'), str(config["outline_level"] - 1))
            ppr.append(outline_lvl)
    
    def _get_alignment(self, align_str: str):
        """转换对齐方式字符串为enum"""
        align_map = {
            "LEFT": WD_PARAGRAPH_ALIGNMENT.LEFT,
            "CENTER": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "RIGHT": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            "JUSTIFY": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
            "DISTRIBUTE": WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE
        }
        return align_map.get(align_str.upper(), WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    def add_cover_page(self, title: str, 
                      metadata: Optional[Dict] = None,
                      ):
        """
        添加封面页
        
        Args:
            title: 主标题
            subtitle: 副标题
            metadata: 元数据字典
            logo_path: Logo图片路径
        """
        # 封面页单独一个节
        section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        
        # # 添加Logo
        # if logo_path and os.path.exists(logo_path):
        #     try:
        #         para = self.doc.add_paragraph()
        #         run = para.add_run()
        #         run.add_picture(logo_path, width=Inches(2))
        #         para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     except:
        #         pass
        
        # 添加主标题
        self.add_heading(title, level=0)
        
        # 添加副标题
        # if subtitle:
        #     self.add_paragraph(subtitle, style="normal", style_override={
        #         "font": {"size": Pt(16), "color": "666666"},
        #         "paragraph": {"alignment": "CENTER", "space_after": Pt(48)}
        #     })
        
        # 添加元数据表格
        if metadata:
            self._add_metadata_table(metadata)
        
        # 添加空行占位
        for _ in range(8):
            self.doc.add_paragraph()
        
        # 添加底部信息
        # bottom_info = self.doc.add_paragraph()
        # bottom_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # run = bottom_info.add_run("智慧园区数字化运营中心")
        # run.font.size = Pt(12)
        # run.font.color.rgb = RGBColor(102, 102, 102)
        
        # 添加分页符
        self.doc.add_page_break()
    
    def _add_metadata_table(self, metadata: Dict):
        """添加元数据表格"""
        if not metadata:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=2)
        table_config = self.styles_config["table"]
        
        # 应用表格样式
        table.style = table_config.get("default_style", "Table Grid")
        table.autofit = False
        # 设置行高
        row_height = table_config.get("row", {}).get("height", Pt(20))
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height
        # 设置列宽
        table.columns[0].width = Inches(1.8)
        table.columns[1].width = Inches(4.2)
        
        # 添加数据
        for key, value in metadata.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
            
            # 样式设置
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = self.doc.styles["Normal"]
        
        # 表格居中
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.doc.add_paragraph()
    
    def add_heading(self, text: str, level: int = 1, 
                   style_override: Optional[Dict] = None):
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别 (1=封面标题, 2=一级标题, 3=二级标题, 4=三级标题)
            style_override: 样式覆盖配置
        """
        if level == 0:
            # 封面标题
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            self._apply_run_style(run, style_override or {})
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return
        if level < 1 or level > 4:
            level = 1
        
        # 强制使用Word内置标题样式（关键！）
        heading_style = f"Heading {level}"
        para = self.doc.add_paragraph(text, style=heading_style)
        
        # 清空默认文本，重新添加（确保样式应用）
        para.clear()
        run = para.add_run(text)
        
        # 应用自定义样式覆盖
        style_key = f"heading{level}"
        if style_key in self.styles_config["styles"]:
            style_config = self.styles_config["styles"][style_key].copy()
            if style_override:
                style_config = self._merge_configs(style_config, style_override)
            
            # 应用字体样式
            self._apply_run_style(run, style_config.get("font", {}))
            
            # 应用段落样式
            para_config = style_config.get("paragraph", {})
            if "alignment" in para_config:
                para.alignment = self._get_alignment(para_config["alignment"])
            
            for key in ["space_before", "space_after"]:
                if key in para_config:
                    setattr(para.paragraph_format, key, para_config[key])
    
    def _apply_run_style(self, run, font_config: Dict):
        """应用运行样式"""
        if "name" in font_config:
            run.font.name = font_config["name"]
            # 中文字体
            r_fonts = run.element.rPr.rFonts
            r_fonts.set(qn("w:eastAsia"), font_config["name"])
            r_fonts.set(qn("w:ascii"), font_config["name"])
        
        if "size" in font_config:
            run.font.size = font_config["size"]
        if "bold" in font_config:
            run.font.bold = font_config["bold"]
        if "italic" in font_config:
            run.font.italic = font_config["italic"]
        if "color" in font_config:
            run.font.color.rgb = RGBColor.from_string(font_config["color"])
        if "underline" in font_config:
            run.font.underline = font_config["underline"]
    
    def add_paragraph(self, text: str, style: str = "normal",
                     style_override: Optional[Dict] = None):
        """
        添加段落
        
        Args:
            text: 段落文本
            style: 样式名称
            style_override: 样式覆盖配置
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        # 应用样式
        if style in self.styles_config["styles"]:
            style_config = self.styles_config["styles"][style].copy()
            if style_override:
                style_config = self._merge_configs(style_config, style_override)
            
            # 字体样式
            self._apply_run_style(run, style_config.get("font", {}))
            
            # 段落样式
            para_config = style_config.get("paragraph", {})
            if "alignment" in para_config:
                para.alignment = self._get_alignment(para_config["alignment"])
            
            for key in ["space_before", "space_after", "left_indent", "right_indent",
                       "first_line_indent", "hanging_indent"]:
                if key in para_config:
                    setattr(para.paragraph_format, key, para_config[key])
            
            if "line_spacing" in para_config:
                para.paragraph_format.line_spacing = para_config["line_spacing"]
    
    def add_table(self, data: List[List[Any]], 
                 headers: Optional[List[str]] = None,
                 title: Optional[str] = None,
                 style: str = "default"):
        """
        添加表格
        
        Args:
            data: 表格数据 (二维列表)
            headers: 表头列表
            title: 表格标题
            style: 表格样式
        """
        if title:
            self.add_paragraph(title, style="normal", style_override={
                "font": {"bold": True,"size":Pt(16)},
                "paragraph": {"space_before": Pt(12), "space_after": Pt(6)}
            })
        
        # 计算表格行列数
        num_rows = len(data)
        num_cols = max(len(row) for row in data) if data else 0
        
        if headers:
            num_cols = max(num_cols, len(headers))
            num_rows += 1
        
        if num_rows == 0 or num_cols == 0:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table_config = self.styles_config["table"]
        
        # 应用表格样式
        table.style = table_config.get("default_style", "Table Grid")
        
        # 设置行高
        row_height = table_config.get("row", {}).get("height", Pt(20))
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = row_height
        
        # 填充表头
        if headers:
            header_config = table_config.get("header", {})
            for i, header_text in enumerate(headers):
                if i < num_cols:
                    cell = table.cell(0, i)
                    cell.text = str(header_text)
                    
                    # 设置表头样式
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                        font_config = header_config.get("font", {})
                        self._apply_run_style(run, font_config)
                        paragraph.alignment = self._get_alignment(header_config.get("alignment", "CENTER"))
                    
                    # 设置背景色
                    bg_color = header_config.get("background_color", "2E74B5")
                    self._set_cell_background(cell, bg_color)
        
        # 填充数据
        start_row = 1 if headers else 0
        row_config = table_config.get("row", {})
        
        for i, row_data in enumerate(data):
            row_idx = start_row + i
            if row_idx >= num_rows:
                continue
                
            for j, cell_data in enumerate(row_data):
                if j >= num_cols:
                    continue
                    
                cell = table.cell(row_idx, j)
                cell.text = str(cell_data)
                
                # 设置数据行样式
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    font_config = row_config.get("font", {})
                    self._apply_run_style(run, font_config)
                    paragraph.alignment = self._get_alignment(row_config.get("alignment", "CENTER"))
                
                # 交替行背景色
                bg_color = row_config.get("background_even" if row_idx % 2 == 0 else "background_odd", "FFFFFF")
                self._set_cell_background(cell, bg_color)
        
        self.doc.add_paragraph()
    
    def _set_cell_background(self, cell, color_hex: str):
        """设置单元格背景色"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading)
        
    def _fix_escape_characters(self, text: str) -> str:
        """修复转义字符问题"""
        # 1. 修复换行符
        text = text.replace('\\n', '\n')  # 处理转义的\n
        
        # 2. 处理其他常见的Markdown换行
        text = text.replace('  \n', '\n')  # Markdown的硬换行
        text = text.replace('<br>', '\n')  # HTML换行
        text = text.replace('<br/>', '\n')
        text = text.replace('<br />', '\n')
        
        # 3. 处理列表项换行
        # 将 "1、xxx\n2、yyy" 转换为正确的换行
        # text = re.sub(r'(\d+、)', r'\n\1', text)  # 1、项前加换行
        # text = re.sub(r'([一二三四五六七八九十]+)、', r'\n\1、', text)  # 中文数字列表项
        # 4. 处理括号编号换行
        text = re.sub(r'([（(]\d+[）)])、', r'\n\1、', text)  # 确保(1)、前有换行
        text = re.sub(r'([（(][a-z][）)])\s*', r'\n    \1 ', text)  # 确保(a)前有换行
        text = re.sub(r'\*\*(\d+\.?\d*(%?|[次起]+)?|数据缺失)\*\*', r'\1', text)  # 确保**124**为124
        
        return text.strip()

    def add_markdown_content(self, markdown_text: str):
        """
        解析并添加Markdown格式的内容
        
        Args:
            markdown_text: Markdown格式的文本
        """
         # 首先处理转义字符
        markdown_text = self._fix_escape_characters(markdown_text)
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # 空行
            if not line:
                self.doc.add_paragraph()
                i += 1
                continue
            
            # 标题检测
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                self.add_heading(text, level=level)
                i += 1
                continue
            
            # 无序列表
            if re.match(r'^([\-\+]|\(\d+\))\s+.+$', line):
                items = []
                while i < len(lines) and re.match(r'^([\-\+]|\(\d+\))\s+.+$', lines[i]):
                    item_text = re.sub(r'^([\-\+]|\(\d+\))\s+', '', lines[i])
                    items.append(item_text)
                    i += 1
                
                for item in items:
                    para = self.doc.add_paragraph()
                    para.style = self.doc.styles["List Bullet"]
                    run = para.add_run(f"·{item}")
                    self._apply_run_style(run, self.styles_config["styles"]["list_item"]["font"])
                continue
             # 数字列表
            # if re.match(r'^(\d|\(\d+\))+、', line):
            #     items = []
            #     while i < len(lines) and re.match(r'^(\d|\(\d+\))+、', lines[i]):
            #         item_text = re.sub(r'^(\d|\(\d+\))+、', '', lines[i])
            #         items.append(item_text)
            #         i += 1
                
            #     for item in items:
            #         para = self.doc.add_paragraph()
            #         para.style = self.doc.styles["List Bullet"]
            #         run = para.add_run(item)
            #         self._apply_run_style(run, self.styles_config["styles"]["list_item"]["font"])
            #     continue
            
            # 表格检测
            if '|' in line and re.match(r'^\|.*\|$', line):
                table_lines = []
                # 收集表格行
                while i < len(lines) and '|' in lines[i] and re.match(r'^\|.*\|$', lines[i]):
                    table_lines.append(lines[i])
                    i += 1
                
                if len(table_lines) >= 2:
                    self._parse_markdown_table(table_lines)
                continue
            
            # 普通段落
            paragraph_text = line
            i += 1
            
            # 合并连续的非空行
            while i < len(lines) and lines[i].strip() and not re.match(r'^[#\*\-\+|]', lines[i]):
                paragraph_text += ' ' + lines[i].strip()
                i += 1
            
            self.add_paragraph(paragraph_text)
    
    def _parse_markdown_table(self, table_lines: List[str]):
        """解析Markdown表格"""
        # 提取表头
        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        
        # 跳过分隔行
        data_start = 2 if len(table_lines) > 1 and '---' in table_lines[1] else 1
        
        # 提取数据
        data = []
        for i in range(data_start, len(table_lines)):
            row = [cell.strip() for cell in table_lines[i].split('|')[1:-1]]
            data.append(row)
        
        # 添加表格
        self.add_table(data, headers)
    
    def add_toc(self):
        """添加目录"""
        if not self.styles_config["toc"]["enabled"]:
            return
        
        # 添加目录标题
        toc_config = self.styles_config["toc"]
        self.add_heading(toc_config["title"], level=2)
        
        # 插入可手动更新的目录占位符（简化版，兼容性更好）
        para = self.doc.add_paragraph()
        # 创建域的XML结构
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        # 将域添加到段落
        para._p.append(fldChar)
        para._p.append(instrText)
        para._p.append(fldChar2)
        para._p.append(fldChar3)
        
        self.doc.add_page_break()

    
    
    def save(self, filepath: str, auto_update_toc: bool = True):
        """保存文档"""
        if self.doc:
            self.doc.save(filepath)
           
            print(f"✅ Word文档已保存: {filepath}")
            return filepath
        return None